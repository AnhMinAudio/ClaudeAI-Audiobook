"""
AnhMin Audio - File Handler
Handle file reading, writing and export to DOCX
"""

import os
import shutil
from pathlib import Path
from typing import Optional, Tuple
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import docx2txt

from config import PROJECTS_DIR


class FileHandler:
    """Handle file operations for the application."""
    
    SUPPORTED_TYPES = {
        '.txt': 'text/plain',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.doc': 'application/msword',
    }
    
    @staticmethod
    def get_project_dir(project_id: int) -> Path:
        """Get the directory for a project's files."""
        project_dir = PROJECTS_DIR / str(project_id)
        project_dir.mkdir(exist_ok=True)
        return project_dir
    
    @staticmethod
    def get_file_type(filepath: str) -> str:
        """Get file type from extension."""
        ext = Path(filepath).suffix.lower()
        return ext.lstrip('.')
    
    @staticmethod
    def get_file_size(filepath: str) -> int:
        """Get file size in bytes."""
        return os.path.getsize(filepath)
    
    @staticmethod
    def format_file_size(size_bytes: int) -> str:
        """Format file size to human readable string."""
        if size_bytes < 1024:
            return f"{size_bytes} B"
        elif size_bytes < 1024 * 1024:
            return f"{size_bytes / 1024:.1f} KB"
        else:
            return f"{size_bytes / (1024 * 1024):.1f} MB"
    
    @classmethod
    def read_file(cls, filepath: str) -> Tuple[str, Optional[str]]:
        """
        Read file content.
        Returns: (content, error_message)
        """
        try:
            path = Path(filepath)
            ext = path.suffix.lower()
            
            if ext == '.txt':
                # Try different encodings
                for encoding in ['utf-8', 'utf-16', 'cp1252', 'latin-1']:
                    try:
                        with open(filepath, 'r', encoding=encoding) as f:
                            return f.read(), None
                    except UnicodeDecodeError:
                        continue
                return "", "Không thể đọc file với các encoding thông dụng"
            
            elif ext in ['.docx', '.doc']:
                content = docx2txt.process(filepath)
                return content, None
            
            else:
                return "", f"Không hỗ trợ định dạng file: {ext}"
                
        except Exception as e:
            return "", f"Lỗi đọc file: {str(e)}"
    
    @classmethod
    def copy_to_project(cls, source_path: str, project_id: int) -> Tuple[str, Optional[str]]:
        """
        Copy a file to project directory.
        Returns: (new_filepath, error_message)
        """
        try:
            source = Path(source_path)
            if not source.exists():
                return "", "File không tồn tại"
            
            project_dir = cls.get_project_dir(project_id)
            
            # Generate unique filename if exists
            dest_name = source.name
            dest_path = project_dir / dest_name
            counter = 1
            while dest_path.exists():
                stem = source.stem
                suffix = source.suffix
                dest_name = f"{stem}_{counter}{suffix}"
                dest_path = project_dir / dest_name
                counter += 1
            
            shutil.copy2(source_path, dest_path)
            return str(dest_path), None
            
        except Exception as e:
            return "", f"Lỗi copy file: {str(e)}"
    
    @classmethod
    def delete_file(cls, filepath: str) -> Optional[str]:
        """
        Delete a file.
        Returns: error_message or None if success
        """
        try:
            path = Path(filepath)
            if path.exists():
                path.unlink()
            return None
        except Exception as e:
            return f"Lỗi xóa file: {str(e)}"
    
    @classmethod
    def export_to_docx(cls, content: str, title: str = "", 
                       output_path: str = None,
                       project_name: str = "",
                       episode_info: str = "") -> Tuple[str, Optional[str]]:
        """
        Export content to DOCX file.
        Returns: (filepath, error_message)
        """
        try:
            doc = Document()
            
            # Set up styles
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(13)
            
            # Add header if project name provided
            if project_name or episode_info:
                header_text = ""
                if project_name:
                    header_text = project_name
                if episode_info:
                    header_text += f" - {episode_info}" if header_text else episode_info
                
                header = doc.add_paragraph()
                header.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = header.add_run(header_text)
                run.bold = True
                run.font.size = Pt(16)
                doc.add_paragraph()  # Empty line
            
            # Add title if provided
            if title:
                title_para = doc.add_paragraph()
                title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = title_para.add_run(title)
                run.bold = True
                run.font.size = Pt(14)
                doc.add_paragraph()  # Empty line
            
            # Add content
            paragraphs = content.split('\n\n')
            for para_text in paragraphs:
                if para_text.strip():
                    para = doc.add_paragraph()
                    para.add_run(para_text.strip())
                    para.paragraph_format.first_line_indent = Inches(0.5)
                    para.paragraph_format.space_after = Pt(6)
            
            # Generate output path if not provided
            if not output_path:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"output_{timestamp}.docx"
                output_path = str(Path.home() / "Downloads" / filename)
            
            # Ensure directory exists
            Path(output_path).parent.mkdir(parents=True, exist_ok=True)
            
            doc.save(output_path)
            return output_path, None
            
        except Exception as e:
            return "", f"Lỗi xuất file DOCX: {str(e)}"
    
    @classmethod
    def export_chat_to_docx(cls, messages: list, project_name: str = "",
                            output_path: str = None) -> Tuple[str, Optional[str]]:
        """
        Export chat history to DOCX file.
        Returns: (filepath, error_message)
        """
        try:
            doc = Document()
            
            # Title
            title = doc.add_paragraph()
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = title.add_run(f"Lịch sử Chat - {project_name}")
            run.bold = True
            run.font.size = Pt(16)
            
            # Date
            date_para = doc.add_paragraph()
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_para.add_run(f"Xuất ngày: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
            doc.add_paragraph()
            
            # Messages
            for msg in messages:
                role = "Bạn" if msg['role'] == 'user' else "Claude"
                
                # Role header
                role_para = doc.add_paragraph()
                role_run = role_para.add_run(f"{role}:")
                role_run.bold = True
                
                # Content
                content_para = doc.add_paragraph()
                content_para.add_run(msg['content'])
                content_para.paragraph_format.left_indent = Inches(0.25)
                
                doc.add_paragraph()  # Separator
            
            # Generate output path
            if not output_path:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"chat_export_{timestamp}.docx"
                output_path = str(Path.home() / "Downloads" / filename)
            
            Path(output_path).parent.mkdir(parents=True, exist_ok=True)
            doc.save(output_path)
            return output_path, None
            
        except Exception as e:
            return "", f"Lỗi xuất chat: {str(e)}"


# Singleton instance
file_handler = FileHandler()
